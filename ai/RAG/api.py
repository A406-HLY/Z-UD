import os
import json
from schemas import SearchResponse, FieldSearchResponse
from keyword_pipeline.pipeline import KeywordExtractionPipeline
from hybrid_retriever import HybridRetriever
from database import save_chunks, load_chunks
current_dir = os.path.dirname(os.path.abspath(__file__))
data_dir = os.path.join(current_dir, 'data')
document_dir = os.path.join(data_dir, 'document')
chunk_dump_dir = os.path.join(data_dir, 'chunks')
template_dir = os.path.join(data_dir, 'template')
concept_dict_path = os.path.join(data_dir, 'concept_dict.json')
original_rules_path = os.path.join(document_dir, 'ssageumjari.txt')
english_to_korean_map_path = os.path.join(data_dir, 'english_to_korean_map.json')
loan_products_path = os.path.join(data_dir, 'loan_products.json')

class AppState:
    grouped_documents = {}
    grouped_doc_vectors = {}
    retrievers = {}
    query_map = {}
    english_to_korean_map = {}
    korean_to_english_map = {}
    evaluation_template = {}
    category_map = {}
    product_map = {}
    query_vectors_cache = {}

    @classmethod
    def load_maps(cls):
        unified_schema_path = os.path.join(template_dir, 'unified_search_schema.json')
        if os.path.exists(unified_schema_path):
            with open(unified_schema_path, 'r', encoding='utf-8') as f:
                schema = json.load(f)
                for item in schema:
                    eng = item['name_en']
                    kor = item['name_ko']
                    cls.english_to_korean_map[eng] = kor
                    cls.korean_to_english_map[kor] = eng
                    cls.query_map[kor] = item['search_query']
                    cls.evaluation_template[kor] = item['default_value']
                    cls.category_map[kor] = item.get('category', '공통')
        if os.path.exists(loan_products_path):
            with open(loan_products_path, 'r', encoding='utf-8') as f:
                cls.product_map = json.load(f)

    @classmethod
    def init_query_vectors(cls):
        """Pre-embeds all static queries from the schema to achieve 0s lookup latency."""
        if not cls.retrievers or not cls.query_map:
            return
        
        # Grab any retriever's model since weights are identical
        if cls.query_vectors_cache:
            return # Already initialized
            
        print('[Startup] Pre-embedding static search queries to memory cache...')
        model = list(cls.retrievers.values())[0].model
        
        queries = []
        kor_keys = []
        for kor_key, query_text in cls.query_map.items():
            queries.append(query_text)
            kor_keys.append(kor_key)
            
        q_vecs = model.encode(queries)
        for i, kor_key in enumerate(kor_keys):
            cls.query_vectors_cache[kor_key] = q_vecs[i]
        print(f'[Startup] Successfully pre-embedded {len(queries)} static search queries.')

AppState.load_maps()

def startup_event():
    print('[Startup] Initializing Vector DB connection...')
    loaded_any = False
    try:
        if not os.path.exists(document_dir):
            os.makedirs(document_dir, exist_ok=True)
        for filename in os.listdir(document_dir):
            if not filename.endswith('.txt'):
                continue
            doc_name = filename.replace('.txt', '')
            try:
                documents, doc_vectors = load_chunks(doc_name)
                if documents and doc_vectors is not None:
                    AppState.grouped_documents[doc_name] = documents
                    AppState.grouped_doc_vectors[doc_name] = doc_vectors
                    AppState.retrievers[doc_name] = HybridRetriever(concept_dict_path=concept_dict_path)
                    print(f'[Startup] Successfully loaded {len(documents)} chunks from Vector DB for {doc_name}.')
                    loaded_any = True
            except Exception as item_e:
                print(f'[Startup Warning] Could not load DB for {doc_name}: {item_e}')
        if not loaded_any:
            print('[Startup] No existing DB rows found for any document. Building from local files...')
            _update_pipeline()
    except Exception as e:
        print(f'[Startup Warning] DB Load failed entirely: {e}. Falling back to normal pipeline...')
        _update_pipeline()

    AppState.init_query_vectors()

def _update_pipeline(raw_text=None):
    total_chunks = 0
    AppState.grouped_documents.clear()
    AppState.grouped_doc_vectors.clear()
    AppState.retrievers.clear()
    if not os.path.exists(document_dir):
        os.makedirs(document_dir, exist_ok=True)
    if not os.path.exists(chunk_dump_dir):
        os.makedirs(chunk_dump_dir, exist_ok=True)
    for filename in os.listdir(document_dir):
        if not (filename.endswith('.txt') or filename.endswith('.docx')):
            continue
        
        file_ext = os.path.splitext(filename)[1].lower()
        doc_name = filename.replace(file_ext, '')
        file_path = os.path.join(document_dir, filename)
        
        doc_text = ""
        if file_ext == '.txt':
            with open(file_path, 'r', encoding='utf-8') as f:
                doc_text = f.read()
        elif file_ext == '.docx':
            try:
                import docx
                doc = docx.Document(file_path)
                doc_text = '\n'.join([para.text for para in doc.paragraphs])
            except ImportError:
                print(f"[Warning] python-docx not installed. Skipping {filename}.")
                continue
                
        print(f'[{doc_name}] Starting Keyword Extraction Pipeline...')
        pipeline = KeywordExtractionPipeline(content=doc_text)
        documents = pipeline.run()
        retriever = HybridRetriever(concept_dict_path=concept_dict_path)
        doc_vectors = retriever.encode_documents(documents)
        AppState.grouped_documents[doc_name] = documents
        AppState.grouped_doc_vectors[doc_name] = doc_vectors
        AppState.retrievers[doc_name] = retriever
        try:
            save_chunks(doc_name, documents, doc_vectors)
            print(f'[{doc_name}] Saved {len(documents)} embedded chunks to DB')
            
            # Save chunks to JSON file for LLM or debugging
            output_json_path = os.path.join(chunk_dump_dir, f"{doc_name}_chunks.json")
            with open(output_json_path, 'w', encoding='utf-8') as json_f:
                import json
                json.dump(documents, json_f, ensure_ascii=False, indent=2)
            print(f'[{doc_name}] Exported chunks to {output_json_path}')
        except Exception as e:
            print(f'[{doc_name} Error] Failed to save DB or JSON: {e}')
        total_chunks += len(documents)
    return total_chunks

def update_single_document(doc_name: str) -> int:
    """Read a single document by name, run keyword pipeline, vector embedding, and update DB/JSON."""
    doc_text = None
    for ext in ['.txt', '.docx']:
        file_path = os.path.join(document_dir, f"{doc_name}{ext}")
        if os.path.exists(file_path):
            if ext == '.txt':
                with open(file_path, 'r', encoding='utf-8') as f:
                    doc_text = f.read()
            elif ext == '.docx':
                try:
                    import docx
                    doc = docx.Document(file_path)
                    doc_text = '\n'.join([para.text for para in doc.paragraphs])
                except ImportError:
                    print(f"[Warning] python-docx not installed. Skipping {doc_name}.docx.")
                    doc_text = ""
            break

    if not doc_text:
        raise FileNotFoundError(f"Document {doc_name} (.txt or .docx) not found or could not be read.")

    print(f'[{doc_name}] Starting Keyword Extraction for single document update...')
    pipeline = KeywordExtractionPipeline(content=doc_text)
    documents = pipeline.run()
    
    retriever = HybridRetriever(concept_dict_path=concept_dict_path)
    doc_vectors = retriever.encode_documents(documents)
    
    # Update active state
    AppState.grouped_documents[doc_name] = documents
    AppState.grouped_doc_vectors[doc_name] = doc_vectors
    AppState.retrievers[doc_name] = retriever

    try:
        save_chunks(doc_name, documents, doc_vectors)
        print(f'[{doc_name}] Saved {len(documents)} embedded chunks to DB')
        
        # Save chunks to JSON file for LLM or debugging
        if not os.path.exists(chunk_dump_dir):
            os.makedirs(chunk_dump_dir, exist_ok=True)
        output_json_path = os.path.join(chunk_dump_dir, f"{doc_name}_chunks.json")
        with open(output_json_path, 'w', encoding='utf-8') as json_f:
            import json
            json.dump(documents, json_f, ensure_ascii=False, indent=2)
        print(f'[{doc_name}] Exported chunks to {output_json_path}')
        
    except Exception as e:
        print(f'[{doc_name} Error] Failed to save DB or JSON: {e}')

    return len(documents)

def update_vector_db():
    num_chunks = _update_pipeline()
    return {'status': 'success', 'message': 'Vector DB updated successfully', 'total_chunks': num_chunks}

def search_rules(user_data: dict):
    if not AppState.retrievers:
        raise RuntimeError('Retrievers not initialized. Call update_vector_db() or startup_event() first.')
    req_uuid = user_data.pop('consultationId', None) or user_data.pop('counselId', None) or user_data.pop('UUID', None) or user_data.pop('uuid', None)
    base_eval_data = {}
    for eng_key, kor_key in AppState.english_to_korean_map.items():
        base_eval_data[eng_key] = {'koreanField': kor_key, 'value': AppState.evaluation_template[kor_key], 'query': AppState.query_map[kor_key], 'matched': []}
    payload = user_data.get('reportInput', user_data.get('user_data', user_data))
    if not isinstance(payload, dict):
        payload = user_data
    enum_translation_map = {'loanPurpose': {'HOME_PURCHASE': '주택 구매', 'REFINANCE': '대환 대출', 'JEONSE_RETURN': '전세보증금 반환', 'LIVING_EXPENSE': '생활 안정 자금'}, 'employmentType': {'EMPLOYEE': '근로자', 'SELF_EMPLOYED': '개인사업자'}}
    for field, mapping in enum_translation_map.items():
        if field in payload and payload[field] in mapping:
            payload[field] = mapping[payload[field]]
    for req_key, val in payload.items():
        if val is not None:
            if req_key in base_eval_data:
                base_eval_data[req_key]['value'] = val
            elif req_key in AppState.korean_to_english_map:
                eng_key = AppState.korean_to_english_map[req_key]
                base_eval_data[eng_key]['value'] = val
    product_results = {doc_name: {} for doc_name in AppState.retrievers.keys()}
    emp_type = payload.get('employmentType', '근로자')
    valid_categories = {'공통', '특수', emp_type}
    filtered_eval_data = {}
    for eng_key, data in base_eval_data.items():
        kor_field = data['koreanField']
        if AppState.category_map.get(kor_field, '공통') in valid_categories:
            filtered_eval_data[eng_key] = data
    for eng_key, data in filtered_eval_data.items():
        field = data['koreanField']
        value = data['value']
        base_query = data['query']
        if value is None:
            for doc_name in AppState.retrievers.keys():
                product_results[doc_name][eng_key] = {'name_ko': field, 'value': None, 'search_query': base_query, 'matched_articles': []}
            continue
        for doc_name, retriever in AppState.retrievers.items():
            doc_documents = AppState.grouped_documents[doc_name]
            doc_vectors = AppState.grouped_doc_vectors[doc_name]
            q_vec_cached = AppState.query_vectors_cache.get(field)
            retrieved = retriever.retrieve(query_text=base_query, target_field=field, documents=doc_documents, doc_vectors=doc_vectors, top_k=None, threshold=0.95, q_vec=q_vec_cached)
            all_matched_chunks = []
            for r in retrieved:
                chunk = r['chunk']
                article_id = chunk.get('article_id', '')
                if article_id:
                    all_matched_chunks.append(article_id)
                else:
                    all_matched_chunks.append(chunk.get('chunk_id', ''))
            seen = set()
            unique_matched = []
            for match in all_matched_chunks:
                if match not in seen:
                    seen.add(match)
                    unique_matched.append(match)
            product_results[doc_name][eng_key] = {'name_ko': field, 'value': value, 'search_query': base_query, 'matched_articles': unique_matched}
    final_result_dict = {}
    for doc_name, fields_dict in product_results.items():
        final_result_dict[doc_name] = {}
        product_meta = AppState.product_map.get(doc_name, {'productName': doc_name})
        if isinstance(product_meta, dict):
            for meta_k, meta_v in product_meta.items():
                final_result_dict[doc_name][meta_k] = meta_v
        else:
            final_result_dict[doc_name]['productName'] = product_meta
        for eng_key, field_data in fields_dict.items():
            final_result_dict[doc_name][eng_key] = FieldSearchResponse(name_ko=field_data['name_ko'], value=field_data['value'], search_query=field_data['search_query'], matched_articles=field_data['matched_articles'])
    return SearchResponse(consultationId=req_uuid, result=final_result_dict)
